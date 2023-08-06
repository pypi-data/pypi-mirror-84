from docx import Document
from docx.shared import Mm
from pandas import DataFrame
import logging
import re


class DocxTemplate:
    """
    The DocxTemplate class makes use of python-docx to modify a word document.
    """

    def __init__(self, filename: str):
        '''constructs the object with the word document. 
        Parameters
        ----------
        filename: file name that points on a valide >2007 MS Word document'''
        self.doc = Document(filename)

    def replaceKeywordByString(self, keyword: str, replacement: str):
        '''replace a keyword by the given replacement string, whatever its position is in the document
        Parameters
        ----------
        keyword: keyword to search for
        replacement: replacement string
        '''
        paragraphs = self._find_paragraphs(keyword)
        self._replaceKeyword(paragraphs, keyword, replacement)

    def replaceKeywordByImages(self, keyword, images, width=None, height=None):
        '''replace a keyword in a paragraph by the given images
        Parameters
        ----------
        keyword: keyword to search for
        images: list of images
        width: expected width of images in the document
        height: expected height of images in the document
        '''
        paragraphs = self._find_paragraphs(keyword)
        self._replaceKeyword(paragraphs, keyword, '')
        for paragraph in paragraphs:
            for image in images:
                run = paragraph.add_run()
                if width and height:
                    run.add_picture(image, width=Mm(
                        width), height=Mm(height))
                elif width:
                    run.add_picture(image, width=Mm(width))
                elif height:
                    run.add_picture(image, height=Mm(height))
                else:
                    run.add_picture(image)

    def findTableByKeyword(self, keyword: str):
        '''find a table in a document by finding a keyword in its cells
        Parameters
        ----------
        keyword: keyword to be found in the table cells
        return the first found table
        '''
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if keyword in cell.text.strip():
                        return table

    def findTableByHeader(self, match: list):
        '''find a table in a document by looking at its first row (header) and by checking if it complies with the colum description
        Parameters
        ----------
        match: colum description = a list describing each column with a regex that should match the column title
        return the first table that complies with the description
        '''

        for table in self.doc.tables:
            found = True
            header = table.rows[0].cells
            if len(match) != len(header):
                continue
            for i in range(0, len(match)):
                if not re.match(match[i], header[i].text.strip(), re.IGNORECASE):
                    found = False
                    break
            if found:
                logging.info('expected table found')
                return table
        logging.info("no table found with the given description")
        return None

    def dropTableContent(self, table):
        '''drop table content 
        Parameters
        ----------
        table: python-docx table to empty'''
        for row in table.rows:
            table._element.remove(row._element)

    def dropTableExceptHeader(self, table):
        '''drop table content except its first row which is called the header
        and return header as a list
        Parameters
        ----------
        table: python-docx table to empty'''
        for row in table.rows[1:]:
            table._element.remove(row._element)

        return [cell.text for cell in table.rows[0].cells]

    def addTableHeader(self, table, header: list):
        '''add a header to the table. This table must be empty.
        Parameters
        ----------
        table: empty python-docx table
        header: list of header names
        '''
        if not table.rows:
            table.add_row().cells
        headercells = table.rows[0].cells
        for colindex, cell in enumerate(headercells):
            cell.text = str(header[colindex])

    def fillTableWithData(self, table, data: DataFrame, from_row: int = 0):
        '''fill the given table with the associated data
        Parameters
        ----------
        table: python-docx table to fill in
        data: pandas Dataframe with the same column number et names than the table to fill in
        from_row: start filling at this given row index (first row is from_row=0)
        '''
        table.autofit = True
        # table.style = 'Table Grid'

        count = from_row
        for _, row in data.iterrows():
            if count <= len(table.rows)-1:
                row_cells = table.rows[count].cells
            else:
                row_cells = table.add_row().cells
            count += 1
            for colindex, cell in enumerate(row_cells):
                if colindex < len(row):
                    cell.text = str(row[colindex])

    def save(self, filename: str):
        '''save the document
        Parameters
        ----------
        filename: file name of the destination file
        '''
        self.doc.save(filename)

    def _find_paragraphs(self, keyword):
        paragraphs = self._findKeywordInSections(keyword)
        paragraphs += self._findKeywordInParagraphs(
            keyword, self.doc.paragraphs)
        paragraphs += self._findKeywordInTables(keyword, self.doc.tables)
        return paragraphs

    def _findKeywordInSections(self, keyword):
        paragraphs = []
        for section in self.doc.sections:
            elements = [section.header, section.footer, section.even_page_header,
                        section.even_page_footer, section.first_page_footer, section.first_page_header]
            for element in elements:
                paragraphs += (self._findKeywordInParagraphs(
                    keyword, element.paragraphs))
                paragraphs += self._findKeywordInTables(
                    keyword, element.tables)
        return paragraphs

    def _findKeywordInParagraphs(self, keyword, paragraphs):
        found = []
        for paragraph in paragraphs:
            if keyword in paragraph.text.strip():
                found.append(paragraph)
        return found

    def _findKeywordInTables(self, keyword, tables):
        paragraphs = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if keyword in paragraph.text.strip():
                            paragraphs.append(paragraph)
        return paragraphs

    def _replaceKeyword(self, paragraphs, keyword, replacement):
        for paragraph in paragraphs:
            paragraph.text = paragraph.text.replace(keyword, replacement)
